import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io

def generate_docx(paper, questions):
    doc = Document()
    
    # Optional: adjust page margins to maximize space
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Master Table to hold everything inside a border box
    master_table = doc.add_table(rows=3, cols=5)
    master_table.style = 'Table Grid'
    master_table.autofit = False
    master_table.allow_autofit = False
    
    # Force table to 100% width
    for w in master_table._tbl.xpath('.//w:tblW'):
        w.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', '5000')
        w.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'pct')
    
    if paper.exam_type == 'university':
        # --- ROW 1: HEADER SECTION ---
        cell_header = master_table.cell(0, 0)
        for col_idx in range(1, 5):
            cell_header.merge(master_table.cell(0, col_idx))
        
        p_c = cell_header.paragraphs[0]
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        r_c = p_c.add_run("DR. BABASAHEB AMBEDKAR TECHNOLOGICAL UNIVERSITY, LONERE\n")
        r_c.bold = True
        r_c.font.size = Pt(12)
        
        exam_title = paper.header_title or "Supplementary Summer Examination - 2024"
        r_c = p_c.add_run(f"{exam_title}\n")
        r_c.bold = True
        r_c.font.size = Pt(11)

        # --- ROW 2: META INFO ---
        cell_meta = master_table.cell(1, 0)
        for col_idx in range(1, 5):
            cell_meta.merge(master_table.cell(1, col_idx))
        
        meta_nested = cell_meta.add_table(rows=3, cols=3)
        
        branch_text = paper.branch or "Artificial Intelligence and Data Science Engineering And Allied"
        meta_nested.cell(0, 0).text = f"Course: B. Tech."
        meta_nested.cell(0, 1).text = f"Branch : {branch_text}"
        meta_nested.cell(0, 2).text = f"Semester : {paper.semester or 'V'}"
        
        meta_nested.cell(1, 0).text = f"Subject Code & Name: {paper.subject.code}  {paper.subject.name}"
        meta_nested.cell(1, 0).merge(meta_nested.cell(1, 2))
        
        meta_nested.cell(2, 0).text = f"Max Marks: 60"
        meta_nested.cell(2, 1).text = f"Date: {paper.exam_date or ''}"
        meta_nested.cell(2, 2).text = f"Duration: 3 Hr."
        
        for r in range(3):
            if len(meta_nested.cell(r, 0).paragraphs[0].runs) > 0:
                meta_nested.cell(r, 0).paragraphs[0].runs[0].font.bold = True

        if len(meta_nested.cell(0, 1).paragraphs[0].runs) > 0:
            meta_nested.cell(0, 1).paragraphs[0].runs[0].font.bold = True
        if len(meta_nested.cell(0, 2).paragraphs[0].runs) > 0:
            meta_nested.cell(0, 2).paragraphs[0].runs[0].font.bold = True
        if len(meta_nested.cell(2, 1).paragraphs[0].runs) > 0:
            meta_nested.cell(2, 1).paragraphs[0].runs[0].font.bold = True
        if len(meta_nested.cell(2, 2).paragraphs[0].runs) > 0:
            meta_nested.cell(2, 2).paragraphs[0].runs[0].font.bold = True
        
        meta_nested.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_nested.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        meta_nested.cell(2, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_nested.cell(2, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # --- ROW 3: INSTRUCTIONS ---
        cell_inst = master_table.cell(2, 0)
        for col_idx in range(1, 5):
            cell_inst.merge(master_table.cell(2, col_idx))
        
        p_inst = cell_inst.paragraphs[0]
        r_inst = p_inst.add_run("Instructions to the Students:\n")
        r_inst.bold = True
        r_inst.italic = True
        p_inst.add_run("1. All the questions are compulsory.\n").italic = True
        p_inst.add_run("2. The level of question/expected answer as per OBE or the Course Outcome (CO) on which the question is based is mentioned in ( ) in front of the question.\n").italic = True
        p_inst.add_run("3. Use of non-programmable scientific calculators is allowed.\n").italic = True
        p_inst.add_run("4. Assume suitable data wherever necessary and mention it clearly.").italic = True

    else:
        # --- ROW 1: HEADER SECTION ---
        # Merge all 5 columns
        cell_header = master_table.cell(0, 0)
        for col_idx in range(1, 5):
            cell_header.merge(master_table.cell(0, col_idx))
        
        # We create a 1-row, 3-col nested table in the header cell for logos and text
        header_nested = cell_header.add_table(rows=1, cols=3)
        # Clear the extra paragraph created by default in the merged cell
        if len(cell_header.paragraphs) > 0 and cell_header.paragraphs[0].text == '':
            p = cell_header.paragraphs[0]
            p._element.getparent().remove(p._element)

        # Logo 1
        cell_l = header_nested.cell(0, 0)
        logo1_path = os.path.join('static', 'logo1.png')
        if os.path.exists(logo1_path):
            p_l = cell_l.paragraphs[0]
            p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_l = p_l.add_run()
            r_l.add_picture(logo1_path, width=Inches(0.8))

        # Center Text
        cell_c = header_nested.cell(0, 1)
        p_c = cell_c.paragraphs[0]
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c = p_c.add_run("SHREEYASH PRATISHTHAN'S\n")
        r_c.font.size = Pt(10)
        
        r_c = p_c.add_run("SHREEYASH COLLEGE OF ENGINEERING & TECHNOLOGY, CHH.SAMBHAJINAGAR\n")
        r_c.bold = True
        r_c.font.size = Pt(12)
        
        r_c = p_c.add_run("Satara Parisar, Beed By-Pass Road, Chh.Sambhajinagar -431010 (M.S.)\n")
        r_c.font.size = Pt(9)
        
        r_c = p_c.add_run("NAAC Accredited, ISO Certified Institute\n")
        r_c.bold = True
        r_c.font.size = Pt(10)
        
        exam_title = paper.header_title or ("MIDSEM EXAM QP" if paper.exam_type == 'class_test' else "UNIVERSITY EXAM QP")
        r_c = p_c.add_run(exam_title)
        r_c.bold = True
        r_c.underline = True
        r_c.font.size = Pt(11)

        # Logo 2
        cell_r = header_nested.cell(0, 2)
        logo2_path = os.path.join('static', 'logo2.jpg')
        if os.path.exists(logo2_path):
            p_r = cell_r.paragraphs[0]
            p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_r = p_r.add_run()
            r_r.add_picture(logo2_path, width=Inches(0.8))

        # Adjust inner table column widths forcefully to stop aggressive text wrapping
        header_nested.autofit = False
        header_nested.allow_autofit = False
        
        cell_l.width = Inches(1.0)
        cell_c.width = Inches(5.0)
        cell_r.width = Inches(1.0)
        
        for row in header_nested.rows:
            row.cells[0].width = Inches(1.0)
            row.cells[1].width = Inches(5.0)
            row.cells[2].width = Inches(1.0)
        

        # --- ROW 2: META INFO ---
        cell_meta = master_table.cell(1, 0)
        for col_idx in range(1, 5):
            cell_meta.merge(master_table.cell(1, col_idx))
        
        meta_nested = cell_meta.add_table(rows=4, cols=2)
        
        meta_nested.cell(0, 0).text = f"Course: Computer Science & Engineering (DS)"
        meta_nested.cell(0, 1).text = f"Sem: {paper.semester or ''}"
        meta_nested.cell(1, 0).text = f"Subject Name: {paper.subject.name}"
        meta_nested.cell(1, 1).text = f"Subject Code: {paper.subject.code}"
        meta_nested.cell(2, 0).text = f"Max Marks: {60 if paper.exam_type == 'university' else 20}"
        meta_nested.cell(2, 1).text = f"Duration: {'3 Hr.' if paper.exam_type == 'university' else '1 Hr.'}"
        meta_nested.cell(3, 0).text = f"Date: {paper.exam_date or ''}"
        
        for r in range(4):
            meta_nested.cell(r, 0).paragraphs[0].runs[0].font.bold = True
            if meta_nested.cell(r, 1).paragraphs[0].runs:
                meta_nested.cell(r, 1).paragraphs[0].runs[0].font.bold = True
            meta_nested.cell(r, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # --- ROW 3: INSTRUCTIONS ---
        cell_inst = master_table.cell(2, 0)
        for col_idx in range(1, 5):
            cell_inst.merge(master_table.cell(2, col_idx))
        
        p_inst = cell_inst.paragraphs[0]
        r_inst = p_inst.add_run("Instructions to the Students:\n")
        r_inst.bold = True
        p_inst.add_run("[1] Use of non-programmable calculator is allowed\n")
        p_inst.add_run("[2] Draw the diagram wherever necessary\n")
        p_inst.add_run("[3] Assume suitable data wherever necessary and mention it clearly")


    # --- ROW 4+: QUESTIONS ---
    # We append directly to the master table
    hdr_cells = master_table.add_row().cells
    hdr_cells[0].text = 'Q. No.'
    hdr_cells[1].text = 'Question'
    hdr_cells[2].text = 'Level'
    hdr_cells[3].text = 'CO'
    hdr_cells[4].text = 'Mark'
    for c in hdr_cells:
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def add_bold_row(lbl, txt, mark):
        row = master_table.add_row().cells
        row[0].text = lbl
        row[1].text = txt
        row[4].text = mark
        for c in row:
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        return row

    def add_q_row(lbl, text, lvl, co):
        row = master_table.add_row().cells
        row[0].text = str(lbl)
        row[1].text = text
        row[2].text = lvl
        row[3].text = co
        for c in [row[0], row[2], row[3], row[4]]:
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return row

    if paper.exam_type == 'university':
        add_bold_row('Q.1', 'Attempt following Questions', '12 X 1')
        for i in range(12):
            if i >= len(questions): break
            q = questions[i]
            q_text = q.question_text
            if q.q_type == 'mcq':
                q_text += f"\n(A) {q.opt_a or ''}      (B) {q.opt_b or ''}\n(C) {q.opt_c or ''}      (D) {q.opt_d or ''}"
            add_q_row(i+1, q_text, 'L1', 'CO1')

        if len(questions) > 12:
            add_bold_row('Q.2', 'Solve Any One of the following', '1 X 6')
            labels = ['(A)', '(B)', '(C)']
            for i in range(2):
                idx = 12 + i
                if idx >= len(questions): break
                add_q_row(labels[i], questions[idx].question_text, 'L2', 'CO2')

        if len(questions) > 14:
            add_bold_row('Q.3', 'Solve Any One of the following', '1 X 6')
            for i in range(2):
                idx = 14 + i
                if idx >= len(questions): break
                add_q_row(labels[i], questions[idx].question_text, 'L2', 'CO2')

        if len(questions) > 16:
            add_bold_row('Q.4', 'Solve Any Two of the following', '2 X 6')
            for i in range(3):
                idx = 16 + i
                if idx >= len(questions): break
                add_q_row(labels[i], questions[idx].question_text, 'L2', 'CO3')

        if len(questions) > 19:
            add_bold_row('Q.5', 'Solve Any Two of the following', '2 X 6')
            for i in range(3):
                idx = 19 + i
                if idx >= len(questions): break
                add_q_row(labels[i], questions[idx].question_text, 'L3', 'CO4')

        if len(questions) > 22:
            add_bold_row('Q.6', 'Solve Any Two of the following', '2 X 6')
            for i in range(3):
                idx = 22 + i
                if idx >= len(questions): break
                add_q_row(labels[i], questions[idx].question_text, 'L3', 'CO4')
    else:
        # Q.1
        add_bold_row('Q.1', 'Attempt following Questions', '1 X 5')
        for i in range(5):
            if i >= len(questions): break
            q = questions[i]
            q_text = q.question_text
            if q.q_type == 'mcq':
                q_text += f"\n(A) {q.opt_a or ''}      (B) {q.opt_b or ''}\n(C) {q.opt_c or ''}      (D) {q.opt_d or ''}"
            add_q_row(i+1, q_text, 'L1' if i < 4 else 'L2', 'CO1')

        # Q.2
        if len(questions) > 5:
            add_bold_row('Q.2', 'Solve Any Two of the following', '2 X 5')
            q2_labels = ['(A)', '(B)', '(C)']
            q2_levels = ['L1', 'L2', 'L2']
            q2_cos = ['CO1', 'CO2', 'CO2']
            for i in range(3):
                idx = 5 + i
                if idx >= len(questions): break
                q = questions[idx]
                add_q_row(q2_labels[i], q.question_text, q2_levels[i], q2_cos[i])
                
        # Q.3
        if len(questions) > 8:
            add_bold_row('Q.3', 'Solve Any One of the following', '1 X 5')
            q3_labels = ['(A)', '(B)']
            for i in range(2):
                idx = 8 + i
                if idx >= len(questions): break
                q = questions[idx]
                add_q_row(q3_labels[i], q.question_text, 'L2', 'CO2')
            
    # Set column widths for main layout
    for row in master_table.rows:
        if len(row.cells) == 5:
            # Note: These values are suggestions; Word will auto-fit tables if not careful
            row.cells[0].width = Inches(0.6)
            row.cells[1].width = Inches(4.5)
            row.cells[2].width = Inches(0.6)
            row.cells[3].width = Inches(0.6)
            row.cells[4].width = Inches(0.8)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
